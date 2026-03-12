"""
Script para convertir archivos Markdown a Word (.docx)
"""

import argparse
import os
import sys
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def markdown_to_word(input_file, output_file=None, base_folder=None, template=None):
    """
    Convierte un archivo Markdown a Word (.docx)
    
    Args:
        input_file: Ruta del archivo Markdown de entrada
        output_file: Ruta del archivo Word de salida (opcional)
        base_folder: Carpeta base para resolver rutas de imágenes (opcional)
        template: Ruta de la plantilla de Word (opcional)
    """
    
    # Si no se especifica archivo de salida, usar el mismo nombre con extensión .docx
    if output_file is None:
        output_file = os.path.splitext(input_file)[0] + '.docx'
    
    # Determinar la carpeta base para las imágenes
    if base_folder is None:
        # Usar la carpeta del archivo markdown como base
        base_folder = os.path.dirname(os.path.abspath(input_file))
        
        # Verificar si las imágenes están en una carpeta IMG sibling (caso docs/ y IMG/ en la raíz)
        # Si el archivo está en docs/, la carpeta IMG estaría en el padre
        parent_folder = os.path.dirname(base_folder)
        img_folder = os.path.join(parent_folder, 'IMG')
        if os.path.exists(img_folder):
            # Actualizar base_folder para buscar en la raíz del proyecto
            base_folder = parent_folder
    
    # Leer el archivo Markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Si hay plantilla, usar proceso especial
    if template and os.path.exists(template):
        # Extraer metadatos y aplicar plantilla
        metadata = extract_metadata_from_markdown(content)
        print(f"Metadatos extraídos: titulo='{metadata['titulo']}', integrantes='{metadata['integrantes'][:50]}...'")
        
        # Crear documento desde plantilla
        doc = Document(template)
        
        # Reemplazar texto en párrafos
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if original_text:
                new_text = original_text
                replaced = False
                for key, value in metadata.items():
                    if value:
                        placeholder = '{{' + key + '}}'
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, value)
                            replaced = True
                if replaced:
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if original_text:
                            new_text = original_text
                            replaced = False
                            for key, value in metadata.items():
                                if value:
                                    placeholder = '{{' + key + '}}'
                                    if placeholder in new_text:
                                        new_text = new_text.replace(placeholder, value)
                                        replaced = True
                            if replaced:
                                for run in paragraph.runs:
                                    run.text = ''
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                                else:
                                    paragraph.add_run(new_text)
        
        # Agregar salto de página después de la portada
        doc.add_page_break()
        
        # Agregar contenido del markdown
        lines = content.split('\n')
        content_started = False
        
        # Cambiar a bucle while para poder procesar tablas (que requieren múltiples líneas)
        i = 0
        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()
            
            # Saltar líneas vacías al inicio
            if not line_stripped:
                if not content_started:
                    i += 1
                    continue
                else:
                    doc.add_paragraph()
                    i += 1
                    continue
            
            # Detectar cuando empieza el contenido real (después de la portada)
            if not content_started:
                if line_stripped.startswith('# '):
                    content_started = True
                else:
                    i += 1
                    continue
            
            if content_started:
                # Título principal
                if line_stripped.startswith('# '):
                    title = doc.add_heading(line_stripped[2:], level=0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    i += 1
                    continue
                
                # Encabezados nivel 1
                if line_stripped.startswith('# '):
                    doc.add_heading(line_stripped[2:], level=1)
                    i += 1
                    continue
                
                # Encabezados nivel 2
                if line_stripped.startswith('## '):
                    doc.add_heading(line_stripped[3:], level=2)
                    i += 1
                    continue
                
                # Encabezados nivel 3
                if line_stripped.startswith('### '):
                    doc.add_heading(line_stripped[4:], level=3)
                    i += 1
                    continue
                
                # Encabezados nivel 4
                if line_stripped.startswith('#### '):
                    doc.add_heading(line_stripped[5:], level=4)
                    i += 1
                    continue
                
                # Tablas - recolectar todas las filas de la tabla
                if line_stripped.startswith('|') and '|' in line_stripped:
                    table_lines = []
                    while i < len(lines) and '|' in lines[i]:
                        table_lines.append(lines[i].strip())
                        i += 1
                    
                    # Procesar la tabla
                    if len(table_lines) >= 2:
                        # Obtener encabezados de la primera fila
                        headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                        
                        # Crear la tabla
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'
                        
                        # Agregar encabezados
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(headers):
                            hdr_cells[j].text = header
                            # Negrita para encabezados
                            for paragraph in hdr_cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        
                        # Agregar filas de datos (saltando la fila de separadores)
                        for row_line in table_lines[2:]:
                            cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                            if len(cells) == len(headers):
                                row_cells = table.add_row().cells
                                for j, cell in enumerate(cells):
                                    row_cells[j].text = cell
                    continue
                
                # Listas con viñetas
                if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                    bullet_text = line_stripped[2:]
                    para = doc.add_paragraph()
                    try:
                        para.style = 'List Bullet'
                    except KeyError:
                        pass
                    add_formatted_text(para, bullet_text)
                    i += 1
                    continue
                
                # Listas numeradas
                if re.match(r'^\d+\.\s', line_stripped):
                    match = re.match(r'^(\d+)\.\s(.*)', line_stripped)
                    if match:
                        bullet_text = match.group(2)
                        para = doc.add_paragraph()
                        try:
                            para.style = 'List Number'
                        except KeyError:
                            pass
                        add_formatted_text(para, bullet_text)
                    i += 1
                    continue
                
                # Imágenes
                img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line_stripped)
                if img_match:
                    image_path = img_match.group(2).lstrip('/')
                    image_full_path = os.path.join(base_folder, image_path)
                    if os.path.exists(image_full_path):
                        try:
                            para = doc.add_paragraph()
                            run = para.add_run()
                            run.add_picture(image_full_path, width=Inches(6))
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"  Advertencia: No se pudo agregar la imagen {image_path}: {e}")
                    i += 1
                    continue
                
                # Texto normal
                para = doc.add_paragraph()
                add_formatted_text(para, line_stripped)
                i += 1
        
        doc.save(output_file)
        print(f"Documento convertido exitosamente con portada: {output_file}")
        return
    
    # Proceso normal sin plantilla
    # Crear el documento Word
    doc = Document()
    
    # Procesar el contenido línea por línea
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Saltar líneas vacías
        if not line:
            i += 1
            continue
        
        # Título principal (# )
        if line.startswith('# '):
            title = doc.add_heading(line[2:], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Encabezados de nivel 1 (#)
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        
        # Encabezados de nivel 2 (##)
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        
        # Encabezados de nivel 3 (###)
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        
        # Encabezados de nivel 4 (####)
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            i += 1
            continue
        
        # Tablas
        if line.startswith('|') and '|' in line:
            # Es una tabla, recolectar todas las filas
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            
            # Procesar la tabla
            if len(table_lines) >= 2:
                # Obtener encabezados de la primera fila
                headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                
                # Crear la tabla
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Agregar encabezados
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    # Negrita para encabezados
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Agregar filas de datos (saltando la fila de separadores)
                for row_line in table_lines[2:]:
                    cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                    if len(cells) == len(headers):
                        row_cells = table.add_row().cells
                        for j, cell in enumerate(cells):
                            row_cells[j].text = cell
            continue
        
        # Listas con viñetas (- )
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:]
            para = doc.add_paragraph()
            try:
                para.style = 'List Bullet'
            except KeyError:
                pass
            add_formatted_text(para, bullet_text)
            i += 1
            continue
        
        # Listas numeradas (1. )
        if re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.*)', line)
            if match:
                bullet_text = match.group(2)
                para = doc.add_paragraph()
                try:
                    para.style = 'List Number'
                except KeyError:
                    pass
                add_formatted_text(para, bullet_text)
                i += 1
                continue
        
        # Imágenes en Markdown ![alt](ruta_imagen)
        image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line)
        if image_match:
            alt_text = image_match.group(1)
            image_path = image_match.group(2)
            
            # Limpiar la ruta de la imagen (quitar / inicial si existe)
            image_path = image_path.lstrip('/')
            
            # Resolver la ruta de la imagen
            if os.path.isabs(image_path):
                image_full_path = image_path
            else:
                # Primero intentar desde la carpeta base (raíz del proyecto)
                image_full_path = os.path.join(base_folder, image_path)
                # Si no existe, intentar desde la carpeta del archivo markdown
                if not os.path.exists(image_full_path):
                    file_dir = os.path.dirname(os.path.abspath(input_file))
                    image_full_path = os.path.join(file_dir, image_path)
            
            # Verificar si la imagen existe
            if os.path.exists(image_full_path):
                try:
                    # Agregar la imagen al documento
                    para = doc.add_paragraph()
                    run = para.add_run()
                    # Usar un ancho máximo de 6 pulgadas para la imagen
                    run.add_picture(image_full_path, width=Inches(6))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  Imagen añadida: {image_path}")
                except Exception as e:
                    print(f"  Advertencia: No se pudo agregar la imagen {image_path}: {e}")
            else:
                print(f"  Advertencia: No se encontró la imagen: {image_full_path}")
            i += 1
            continue
        
        # Texto normal
        para = doc.add_paragraph()
        add_formatted_text(para, line)
        i += 1
    
    # Guardar el documento
    doc.save(output_file)
    print(f"Documento convertido exitosamente: {output_file}")


def add_formatted_text(paragraph, text):
    """
    Agrega texto con formato (nebrreteta, cursiva) al párrafo
    """
    # Buscar texto en nebrreteta **texto**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Texto en nebrreteta
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Texto en cursiva
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif '_' in part and not part.startswith('__'):
            # Texto en cursiva con underscore
            run = paragraph.add_run(part.replace('_', ''))
            run.font.italic = True
        else:
            # Texto normal
            paragraph.add_run(part)


def extract_metadata_from_markdown(content):
    """
    Extrae metadatos del contenido Markdown
    
    Args:
        content: Contenido del archivo Markdown
    
    Returns:
        Dictionary con los metadatos extraídos
    """
    metadata = {
        'titulo': '',
        'curso': '',
        'profesor': '',
        'periodo': '',
        'fecha': '',
        'integrantes': '',
        'escuela': 'Escuela de Transformación Digital'
    }
    
    lines = content.split('\n')
    
    # Buscar el título principal (# )
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            metadata['titulo'] = line[2:].strip()
            break
    
    # Buscar campos en el contenido (formato: **campo:** valor)
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Buscar curso
        if '**Curso:**' in line or '**curso:**' in line:
            match = re.search(r'\*\*(?:Curso|curso):\*\*\s*(.+)', line)
            if match:
                metadata['curso'] = match.group(1).strip()
        
        # Buscar docente/profesor
        if '**Docente:**' in line or '**Docente:**' in line or '**profesor:**' in line:
            match = re.search(r'\*\*(?:Docente|profesor):\*\*\s*(.+)', line)
            if match:
                metadata['profesor'] = match.group(1).strip()
        
        # Buscar período
        if '**Período Académico:**' in line or '**Período:**' in line or '**periodo:**' in line:
            match = re.search(r'\*\*(?:Período Académico|Período|periodo):\*\*\s*(.+)', line)
            if match:
                metadata['periodo'] = match.group(1).strip()
        
        # Buscar fecha de entrega
        if '**Fecha de Entrega:**' in line or '**Fecha:**' in line or '**fecha:**' in line:
            match = re.search(r'\*\*(?:Fecha de Entrega|Fecha|fecha):\*\*\s*(.+)', line)
            if match:
                metadata['fecha'] = match.group(1).strip()
        
        # Buscar estudiantes (lista)
        if '**Estudiantes:**' in line or '**estudiantes:**' in line:
            # Recolectar la lista de estudiantes
            students = []
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if next_line.startswith('- ') or next_line.startswith('* '):
                    students.append(next_line[2:].strip())
                    i += 1
                elif next_line == '':
                    i += 1
                    continue
                else:
                    break
            metadata['integrantes'] = '\n'.join(students)
            continue
        
        i += 1
    
    return metadata


def convert_all_docs(docs_folder='docs', template=None):
    """
    Convierte todos los archivos Markdown de la carpeta docs a Word
    
    Args:
        docs_folder: Ruta de la carpeta con archivos markdown
        template: Ruta de la plantilla de Word (opcional)
    """
    if not os.path.exists(docs_folder):
        print(f"Error: No se encontró la carpeta {docs_folder}")
        return
    
    # Buscar todos los archivos .md en la carpeta
    md_files = [f for f in os.listdir(docs_folder) if f.endswith('.md')]
    
    if not md_files:
        print(f"No se encontraron archivos .md en {docs_folder}")
        return
    
    print(f"Se encontraron {len(md_files)} archivo(s) .md en {docs_folder}")
    print("=" * 50)
    
    for md_file in md_files:
        input_path = os.path.join(docs_folder, md_file)
        output_path = os.path.splitext(input_path)[0] + '.docx'
        print(f"Convirtiendo: {md_file} -> {os.path.basename(output_path)}")
        
        try:
            markdown_to_word(input_path, output_path, template=template)
        except Exception as e:
            print(f"Error al convertir {md_file}: {e}")
        print()
    
    print("=" * 50)
    print(f"Conversión completada: {len(md_files)} archivo(s) convertido(s)")


def main():
    """Función principal"""
    # Configurar argparse
    parser = argparse.ArgumentParser(
        description='Conversor de Markdown a Word (.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos de uso:
  python markdown_to_word.py                           # Convierte docs/Modelado.md
  python markdown_to_word.py archivo.md                  # Convierte un archivo específico
  python markdown_to_word.py archivo.md resultado.docx   # Especifica archivo de salida
  python markdown_to_word.py --all                       # Convierte todos los .md de docs/
  python markdown_to_word.py --all -t plantilla.docx    # Convierte todos con plantilla
  python markdown_to_word.py archivo.md -t plantilla.docx  # Convierte con plantilla
        """
    )
    
    parser.add_argument(
        'input',
        nargs='?',
        default=None,
        help='Ruta del archivo Markdown a convertir (opcional)'
    )
    
    parser.add_argument(
        'output',
        nargs='?',
        default=None,
        help='Ruta del archivo Word de salida (opcional)'
    )
    
    parser.add_argument(
        '--all', '-a',
        action='store_true',
        dest='convert_all',
        help='Convertir todos los archivos .md de la carpeta docs'
    )
    
    parser.add_argument(
        '--docs-folder', '-d',
        type=str,
        default='docs',
        help='Carpeta donde buscar archivos markdown (default: docs)'
    )
    
    parser.add_argument(
        '--template', '-t',
        type=str,
        default=None,
        help='Ruta de la plantilla de Word con campos {{titulo}}, {{integrantes}}, etc.'
    )
    
    args = parser.parse_args()
    
    print("=" * 50)
    print("Conversor de Markdown a Word (.docx)")
    print("=" * 50)
    print()
    
    # Verificar argumentos
    if args.convert_all:
        # Convertir todos los archivos de la carpeta docs
        convert_all_docs(args.docs_folder, args.template)
    elif args.input:
        # Convertir archivo específico
        input_file = args.input
        output_file = args.output
        
        # Verificar si existe
        if not os.path.exists(input_file):
            print(f"Error: No se encontró el archivo {input_file}")
            print("\nUso: python markdown_to_word.py archivo.md [archivo.docx]")
            sys.exit(1)
        
        # Ejecutar conversión
        try:
            markdown_to_word(input_file, output_file, template=args.template)
        except Exception as e:
            print(f"Error durante la conversión: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)
    else:
        # Usar el archivo del proyecto por defecto
        input_file = os.path.join(args.docs_folder, "Modelado.md")
        
        # Verificar si existe
        if not os.path.exists(input_file):
            print(f"Error: No se encontró el archivo {input_file}")
            print(f"\nUso: python markdown_to_word.py archivo.md [archivo.docx]")
            print(f"     python markdown_to_word.py --all")
            sys.exit(1)
        
        output_file = None
        
        # Ejecutar conversión
        try:
            markdown_to_word(input_file, output_file, template=args.template)
        except Exception as e:
            print(f"Error durante la conversión: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)


if __name__ == "__main__":
    main()
